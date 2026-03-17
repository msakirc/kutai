# documents.py — document text extraction

import os
from src.infra.logging_config import get_logger
logger = get_logger("tools.documents")

async def read_pdf(filepath: str, max_pages: int = 20) -> str:
    """Extract text from a PDF file."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(filepath)
        total = len(doc)
        pages = doc[:min(max_pages, total)]
        text = "\n\n".join(p.get_text() for p in pages)
        suffix = f"\n\n[... {total - max_pages} more pages not shown]" if total > max_pages else ""
        logger.info("read pdf", filepath=filepath, pages=min(max_pages, total), total=total)
        return text + suffix
    except Exception as e:
        logger.error("pdf read failed", filepath=filepath, error=str(e))
        return f"Error reading PDF: {e}"

async def read_docx(filepath: str) -> str:
    """Extract text from a .docx file."""
    try:
        from docx import Document
        doc = Document(filepath)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        logger.info("read docx", filepath=filepath, paragraphs=len(doc.paragraphs))
        return text
    except Exception as e:
        logger.error("docx read failed", filepath=filepath, error=str(e))
        return f"Error reading DOCX: {e}"

async def read_spreadsheet(filepath: str, sheet: str | None = None, max_rows: int = 200) -> str:
    """Extract content from an Excel or CSV file."""
    try:
        import openpyxl, csv
        if filepath.endswith(".csv"):
            with open(filepath, newline="", encoding="utf-8", errors="replace") as f:
                rows = list(csv.reader(f))[:max_rows]
            return "\n".join(",".join(row) for row in rows)
        wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
        ws = wb[sheet] if sheet and sheet in wb.sheetnames else wb.active
        rows = []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i >= max_rows:
                break
            rows.append("\t".join("" if v is None else str(v) for v in row))
        logger.info("read spreadsheet", filepath=filepath, rows=len(rows))
        return "\n".join(rows)
    except Exception as e:
        logger.error("spreadsheet read failed", filepath=filepath, error=str(e))
        return f"Error reading spreadsheet: {e}"

async def extract_text(filepath: str) -> str:
    """Auto-detect file type and extract text."""
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        return await read_pdf(filepath)
    elif ext == ".docx":
        return await read_docx(filepath)
    elif ext in (".xlsx", ".xls", ".xlsm", ".csv", ".tsv"):
        return await read_spreadsheet(filepath)
    elif ext in (".txt", ".md", ".rst", ".log", ".json", ".yaml", ".yml", ".toml"):
        with open(filepath, encoding="utf-8", errors="replace") as f:
            return f.read()
    else:
        return f"Unsupported file type: {ext}"
